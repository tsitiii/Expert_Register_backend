import pymupdf
from docx import Document
from io import BytesIO
import re
from typing import Dict, List, Optional, Union
from collections import OrderedDict

class ResumeParser:
    def __init__(self, file_content: bytes, file_extension: str):
        self.file_content = file_content
        self.file_extension = file_extension.lower()
        self.parsed_data = OrderedDict()  # Preserve section order
        self.current_section = None

    def parse(self) -> Dict:
        try:
            if self.file_extension == 'pdf':
                self._parse_pdf()
            elif self.file_extension in ['docx', 'doc']:
                self._parse_docx()
            else:
                raise ValueError("Unsupported file format")
            
            return dict(self.parsed_data)
        except Exception as e:
            raise ValueError(f"Error parsing resume: {str(e)}")

    def _parse_pdf(self):
        with pymupdf.open(stream=BytesIO(self.file_content), filetype='pdf') as doc:
            full_text = ""
            for page in doc:
                full_text += page.get_text()
            self._parse_text_content(full_text)
            
            # PDF table extraction
            for page_num, page in enumerate(doc):
                tables = page.find_tables()
                if tables:
                    for table_num, table in enumerate(tables):
                        table_data = table.extract()
                        table_key = f"table_{page_num}_{table_num}"
                        self._add_table(table_key, table_data)

    def _parse_docx(self):
        with BytesIO(self.file_content) as file_buffer:
            doc = Document(file_buffer)
            
            # Process all paragraphs and tables in document order
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    para = self._get_paragraph_text(element)
                    self._process_paragraph(para)
                elif element.tag.endswith('tbl'):  # Table
                    table = self._get_table_data(element)
                    table_key = f"{self.current_section}_table" if self.current_section else "table"
                    self._add_table(table_key, table)

    def _get_paragraph_text(self, element):
        """Extract text from a paragraph element"""
        text = []
        for run in element.xpath(".//w:t"):
            text.append(run.text)
        return ''.join(text).strip()

    def _get_table_data(self, element):
        """Extract data from a table element"""
        table_data = []
        for row in element.xpath(".//w:tr"):
            row_data = []
            for cell in row.xpath(".//w:tc"):
                cell_text = []
                for para in cell.xpath(".//w:p"):
                    cell_text.append(self._get_paragraph_text(para))
                row_data.append(' '.join(cell_text).strip())
            table_data.append(row_data)
        return table_data

    def _process_paragraph(self, text: str):
        """Process a paragraph of text"""
        if not text:
            return
            
        # Detect section headers (text ending with colon)
        if text.endswith(':') and len(text) < 50:
            self.current_section = text[:-1].lower().replace(' ', '_')
            self.parsed_data[self.current_section] = {
                'text': [],
                'tables': OrderedDict()
            }
            return
            
        # Add text to current section
        if self.current_section:
            self.parsed_data[self.current_section]['text'].append(text)
        else:
            if 'header' not in self.parsed_data:
                self.parsed_data['header'] = {
                    'text': [],
                    'tables': OrderedDict()
                }
            self.parsed_data['header']['text'].append(text)

    def _add_table(self, table_key: str, table_data: List[List[str]]):
        """Add a table to the appropriate section"""
        if not table_data:
            return
            
        # Find appropriate section
        section = self.current_section if self.current_section else 'header'
        
        if section not in self.parsed_data:
            self.parsed_data[section] = {
                'text': [],
                'tables': OrderedDict()
            }
            
        # Ensure table key is unique
        base_key = table_key
        counter = 1
        while table_key in self.parsed_data[section]['tables']:
            table_key = f"{base_key}_{counter}"
            counter += 1
            
        self.parsed_data[section]['tables'][table_key] = table_data

    def _parse_text_content(self, text: str):
        """Parse text content (for PDFs)"""
        lines = text.split('\n')
        for line in lines:
            self._process_paragraph(line)


def read_uploaded_file(file_object, file_extension: str) -> dict:
    """Parse resume/CV file and return structured data with tables"""
    file_content = file_object.read()
    
    try:
        parser = ResumeParser(file_content, file_extension)
        raw_data = parser.parse()
        
        # Post-process to structure the data
        structured_data = {}
        for section, content in raw_data.items():
            structured_data[section] = {
                'text': '\n'.join(content['text']),
                'tables': dict(content['tables'])
            }
        
        return structured_data
    except Exception as e:
        raise ValueError(f"Failed to parse resume: {str(e)}")